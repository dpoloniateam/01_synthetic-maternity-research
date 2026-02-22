import docx
from docx.shared import Pt

def add_heading(target_p, text, level):
    p = target_p.insert_paragraph_before(text)
    # We'll just make it bold if we can't reliably guess the heading style name
    p.runs[0].bold = True
    return p

def add_paragraph(target_p, text):
    return target_p.insert_paragraph_before(text)

doc = docx.Document("docs/Repository/Paper Designing Sharper User Research_v0.docx")

insert_index = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "5. Results":
        insert_index = i
        break

if insert_index != -1:
    # Insert before the next paragraph after "5. Results"
    # Actually there are several empty paragraphs after "5. Results". Let's insert before insert_index + 1
    target_p = doc.paragraphs[insert_index + 1]
    
    h1 = add_heading(target_p, "5.1 Evaluative Analysis of Conversational Yield: From Surface Needs to Latent Insights", 2)
    
    add_paragraph(target_p, "A fundamental challenge in the Front End of Innovation (FEI) is transcending surface-level user requests to uncover the latent, structural, and emotional drivers that inform breakthrough service design. To evaluate the efficacy of the Synthetic Design Laboratory as a micro-sensing capability, we conducted a comparative analysis of the conversational data yielded by two distinct prompt calibrations of the Generative Interviewer Agent.")
    
    add_paragraph(target_p, "In accordance with the knowledge-based view (KBV), the iterative calibration of the interviewer agent is framed as a knowledge-creation routine. This routine functions to convert tacit user vulnerabilities and complex social realities into explicit innovation requirements before human fieldwork begins, thereby systematically augmenting the organization's absorptive capacity.")
    
    p = add_paragraph(target_p, "Our analysis utilized a structured classification rubric that delineates between Needs (i.e., explicit, logistical, or transactional requirements focusing on \"what\" the user wants) and Insights (i.e., deeper understandings revealing hidden truths, structural tensions, and underlying reasons for user actions, focusing on \"why\" the user behaves or feels a certain way).")
    
    # Needs some bolding for keywords
    p.runs[0].text = "Our analysis utilized a structured classification rubric that delineates between "
    p.add_run("Needs").bold = True
    p.add_run(" (i.e., explicit, logistical, or transactional requirements focusing on \"what\" the user wants) and ")
    p.add_run("Insights").bold = True
    p.add_run(" (i.e., deeper understandings revealing hidden truths, structural tensions, and underlying reasons for user actions, focusing on \"why\" the user behaves or feels a certain way).")

    add_heading(target_p, "5.1.1 Baseline Extraction (Run 1): The \"Customer Service\" Paradigm", 3)
    add_paragraph(target_p, "In the initial simulation, the interviewer agent operated with a generalised objective to identify support gaps. However, lacking strict qualitative constraints, the Large Language Model (LLM) defaulted to a solution-oriented \"customer service\" paradigm. The extracted data yielded 100% surface-level Needs, primarily capturing logistical requests (e.g., formal excuse letters). While actionable, this functional data failed to uncover the profound systemic friction causing the participant's distress.")

    add_heading(target_p, "5.1.2 Calibrated Extraction (Mark V2): The Qualitative Researcher Paradigm", 3)
    add_paragraph(target_p, "Following iterative calibration, the agent was strictly constrained to investigate emotional states and structural barriers, adopting the posture of an empathetic ethnographic researcher. This calibration profoundly shifted the data yield, producing transcripts composed of 75% Insights and only 25% Needs. This dramatic enhancement reveals the capacity of the Synthetic Design Laboratory to operate as a high-fidelity micro-sensing capability at the FEI. By simulating real-world tensions and power dynamics, the laboratory significantly enhances the ecological value of the research design upstream.")
    
    add_paragraph(target_p, "The calibrated synthetic user generated deep, reflective elaboration, yielding powerful latent insights critical for human-centered design in maternity care:")

    bullets = [
        ("The \"Stigma of Accommodation\" Insight:", " The synthetic participant revealed that requesting logistical support—such as a stool for physical relief at work or assignment extensions at university—is paradoxically perceived as a capitulation. Within environments characterized by institutional skepticism, utilizing accommodations is internalized not as a right, but as a \"demographic failure\" that merely serves to \"prove them right\" regarding her perceived incapacity."),
        ("The \"Private Struggle\" Insight:", " The data surfaced a paralyzing latent belief that \"needing help equals failure.\" Having internalized the judgment of family and peers, the participant equated asking for support with confirming negative societal stereotypes. Consequently, vulnerable users routinely isolate themselves, hiding severe pain and food insecurity as a behavioral response to mitigate the risk of institutional let-down."),
        ("The \"Competent Patient\" Paradox:", " The transcripts additionally revealed a critical tension wherein the user's attempts to navigate complex care pathways assertively are met with systemic resistance. Being \"too knowledgeable\" or advocating too strongly creates unexpected friction with authority figures in both healthcare and academic settings, further marginalizing the patient.")
    ]
    
    for bold_text, normal_text in bullets:
        bp = target_p.insert_paragraph_before("- ")
        bp.add_run(bold_text).bold = True
        bp.add_run(normal_text)

    add_heading(target_p, "5.1.3 Governance and Replicability", 3)
    add_paragraph(target_p, "To ensure the responsible deployment of this knowledge-creation routine, robust governance mechanisms were instrumental. The continual generation of an \"Audit Trail\"—encompassing granular JSON logs of exact prompt-reply payloads and dynamic summary reports of token utilization—ensures ethical transparency and full replicability of the synthetic user research. This pipeline clarifies the methodological boundary: synthetic data is operationalized strictly as a design testbed to sharpen research instruments, rather than as a substitute for real user voices.")

    doc.save("docs/Repository/Paper Designing Sharper User Research_v1.docx")
    print("Document successfully updated and saved as Paper Designing Sharper User Research_v1.docx")
else:
    print("Could not find '5. Results' section.")
