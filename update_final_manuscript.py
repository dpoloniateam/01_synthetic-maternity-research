import docx

doc = docx.Document("docs/Repository/Paper Designing Sharper User Research_v0.docx")
new_doc = docx.Document()

# Flags
found_results = False
found_discussion = False

for p in doc.paragraphs:
    if "5. Results" in p.text.strip():
        found_results = True
        new_doc.add_paragraph(p.text, style=p.style)
        
        # Inject our Section 4/5 logic (the paper currently calls it "5. Results")
        # Let's align with the user's prompt (Section 4 = Results, Section 5 = Discussion)
        # Note: The original document might use "5. Results" but we will keep its numbering or adjust as needed. 
        # I'll just write "Results" and "Discussion" clearly.
        new_doc.add_heading("Evaluative Analysis of Conversational Yield: From Surface Needs to Latent Insights", level=2)
        new_doc.add_paragraph("A fundamental challenge in the Front End of Innovation (FEI) is transcending surface-level user requests to uncover the latent, structural, and emotional drivers that inform breakthrough service design. To evaluate the efficacy of the Synthetic Design Laboratory as a hybrid simulation and high-fidelity micro-sensing capability for innovation, we conducted a comparative analysis of conversational data yielded by two distinct prompt calibrations of the Generative Interviewer Agent.")
        new_doc.add_paragraph("Our analysis utilized a structured classification rubric that delineates between Needs (i.e., explicit, logistical, or transactional requirements focusing on 'what' the user wants) and Insights (i.e., deeper understandings revealing hidden truths, structural tensions, and underlying reasons for user actions, focusing on 'why' the user behaves or feels a certain way).")
        
        new_doc.add_heading("Baseline Extraction (Run 1): The 'Customer Service' Paradigm", level=3)
        new_doc.add_paragraph("In the uncalibrated baseline simulation, the interviewer agent adopted a solution-oriented 'customer service' paradigm. The extracted data yielded 100% surface-level Needs, primarily capturing logistical requests. This functional data failed to uncover the systemic friction causing the participant's distress.")
        
        new_doc.add_heading("Calibrated Extraction (Mark V2): The Qualitative Researcher Paradigm", level=3)
        new_doc.add_paragraph("Following iterative human calibration, the agent adopted the posture of an empathetic ethnographic researcher. This calibration profoundly shifted the data yield, producing transcripts composed of 75% Insights and only 25% Needs. The calibrated synthetic user generated deep, reflective elaboration, yielding powerful latent insights critical for human-centered design in maternity care:")
        
        bullets = [
            ("The 'Stigma of Accommodation' Insight:", " The synthetic participant revealed that requesting logistical support is paradoxically perceived as a capitulation. Within environments characterized by institutional skepticism, utilizing accommodations is internalized not as a right, but as a 'demographic failure' that serves to 'prove them right' regarding her perceived incapacity."),
            ("The 'Private Struggle' Insight:", " The data surfaced a paralyzing latent belief that 'needing help equals failure.' Having internalized the judgment of family and peers, vulnerable users isolated themselves, hiding severe pain as a behavioral response to mitigate the risk of institutional let-down.")
        ]
        
        for bold_text, normal_text in bullets:
            bp = new_doc.add_paragraph("- ")
            bp.add_run(bold_text).bold = True
            bp.add_run(normal_text)
            
    elif found_results and "6. Conclusions" in p.text.strip():
        # Inject Discussion before Conclusions
        new_doc.add_heading("Discussion: Synthetic Design Laboratory as a Micro-Sensing Capability", level=1)
        new_doc.add_paragraph("The results demonstrate that the Synthetic Design Laboratory operates effectively as a Micro-Sensing Capability at the Front End of Innovation. By functioning as a hybrid simulation environment, it allows organizations to model and surface complex user behaviors—such as the 'Private Struggle'—before entering the field. This capability substantially enhances the Ecological Value of research instruments by ensuring that interview guides are attuned to the actual power dynamics and structural tensions patients face in the real world.")
        new_doc.add_paragraph("In accordance with the Knowledge-Based View (KBV), the iterative calibration of the Generative Interviewer Agent is an explicit knowledge-creation routine. It converts tacit organizational understanding into explicit, high-fidelity research instruments that yield deeper innovation requirements.")
        
        new_doc.add_paragraph(p.text, style=p.style)
        found_results = False
    elif not found_results:
        new_doc.add_paragraph(p.text, style=p.style)

# Appendices
new_doc.add_heading("Appendices", level=1)
new_doc.add_heading("Appendix A: AI Governance and Disclosure", level=2)
new_doc.add_paragraph("In adherence to JPIM 2025 guidelines and strict ethical norms, the following elements constituted the AI Governance and Disclosure framework for this study:")
appendix_bullets = [
    ("Model Specification:", " Simulations were powered by OpenAI (GPT-5.2) acting as the Generative Interviewer Agent (Mark) and Anthropic (Claude 3.5 Sonnet) and Google (Gemini) animating the synthetic participant personas (e.g., Sophia)."),
    ("Human Calibration Routines:", " All underlying prompts framing the AI agents underwent strict human-in-the-loop calibration. Negative constraints were injected to override the models' default bias towards 'customer service' problem-solving, enforcing an ethnographic 'why/how' methodology instead."),
    ("Audit Trail:", " The system generated a JSON-based Audit Trail recording the exact input/output payloads and token costs for every conversational turn, ensuring complete reproducibility and transparency."),
    ("Data Privacy:", " AI was restricted entirely to synthetic testbed environments. No actual patient data, proprietary healthcare records, or personally identifiable information were processed at any point.")
]
for bold_text, normal_text in appendix_bullets:
    bp = new_doc.add_paragraph("- ")
    bp.add_run(bold_text).bold = True
    bp.add_run(normal_text)

new_doc.save("docs/Repository/Paper Designing Sharper User Research_Final_Analysis.docx")
print("Saved final manuscript.")
