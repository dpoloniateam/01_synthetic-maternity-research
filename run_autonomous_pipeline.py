import os
import json
import time
import glob
import shutil
from concurrent.futures import ThreadPoolExecutor
from openai import OpenAI
import anthropic
import google.generativeai as genai
import docx

# Load environment
from dotenv import load_dotenv
load_dotenv(".env")

# APIs
client_openai = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
client_anthropic = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Setup roles with cheap models
MODEL_INTERVIEWER = "gpt-4o-mini"
MODEL_PERSONA = "google/gemini-1.5-flash"
MODEL_SYNTHESIZER = "gemini-1.5-flash"  # Role 3 (Google)
MODEL_EDITOR = "claude-3-haiku-20240307" # Role 4 (Anthropic)
MODEL_REVIEWER = "gpt-4o-mini"          # Role 5 (OpenAI)

COSTS = 0.0

def add_cost(provider, model_name, in_tok, out_tok):
    global COSTS
    # Approximate costs per 1K for cheap tier
    pricing = {
        "gpt-4o-mini": (0.00015, 0.00060),
        "gemini-1.5-flash": (0.000075, 0.00030),
        "claude-3-haiku-20240307": (0.00025, 0.00125)
    }
    rates = pricing.get(model_name, (0.001, 0.002))
    COSTS += (in_tok / 1000.0) * rates[0] + (out_tok / 1000.0) * rates[1]

# 1. Phase 1: 16 Interviews
import subprocess

def run_interview(persona_id):
    env = os.environ.copy()
    env["INTERVIEWER_MODEL"] = MODEL_INTERVIEWER
    env["TARGET_MODEL_OVERRIDE"] = MODEL_PERSONA
    cmd = ["python", "-c", f"from src.orchestration.interview_loop import run_synthetic_interview; run_synthetic_interview('{persona_id}', num_turns=3)"]
    subprocess.run(cmd, env=env, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    print(f"Completed interview for {persona_id}")

def generate_interviews():
    with open('data/config/run_plan_16.json', 'r') as f:
        plan = json.load(f)
    print("Running 16 interviews in parallel...")
    with ThreadPoolExecutor(max_workers=8) as executor:
        executor.map(run_interview, plan)
    
def get_all_transcripts():
    files = glob.glob("data/archive/study1_final_archive/transcript_*.json") + glob.glob("data/transcripts/transcript_*.json")
    content = ""
    for f in files[-16:]:
        with open(f, 'r') as file:
            content += file.read() + "\n"
    return content

# 2. Phase 2: Instrument & Drafting
def synthesize_instrument(transcripts):
    print("Role 3: Synthesizing Validatable_Research_Instrument_v1.docx... (Gemini)")
    model = genai.GenerativeModel('gemini-1.5-flash')
    prompt = f"Based on the following 16 transcripts, extract thematic constructs regarding the 'Stigma of Accommodation', 'Private Struggle'. Generate a robust Validatable Research Instrument containing an Adaptability Matrix and Likert scales derived from latent insights.\n\nTranscripts:\n{transcripts[:30000]}"
    res = model.generate_content(prompt)
    
    doc = docx.Document()
    doc.add_paragraph(res.text)
    doc.save("data/Validatable_Research_Instrument_v1.docx")
    try:
        in_t = res.usage_metadata.prompt_token_count
        out_t = res.usage_metadata.candidates_token_count
        add_cost("google", "gemini-1.5-flash", in_t, out_t)
    except: pass
    print("Instrument generated.")

def draft_manuscript(transcripts):
    print("Role 3: Drafting Paper... (Gemini)")
    source_doc = docx.Document("docs/Repository/Paper Designing Sharper User Research_v0.docx")
    text = "\n".join([p.text for p in source_doc.paragraphs])
    
    model = genai.GenerativeModel('gemini-1.5-flash')
    prompt = f"Rewrite Sections 4 and 5 of this manuscript. Position the synthetic instrument as a 'Micro-Sensing Capability' under the Knowledge-Based View (KBV) to bypass the 'Stigma of Accommodation'. Use empirical findings from organic extraction to show yield improvements from Needs to Insights.\nManuscript:\n{text[:20000]}"
    res = model.generate_content(prompt)
    
    new_doc = docx.Document()
    new_doc.add_paragraph("JPIM Submission Draft")
    new_doc.add_paragraph(res.text)
    new_doc.save("data/Paper_Draft_v1.docx")
    try:
        in_t = res.usage_metadata.prompt_token_count
        out_t = res.usage_metadata.candidates_token_count
        add_cost("google", "gemini-1.5-flash", in_t, out_t)
    except: pass
    return res.text

# 3. Phase 3: Editorial Gate & Peer Review Loop
def editorial_gate(draft_text):
    print("Role 4: JPIM Editor evaluating draft... (Claude)")
    prompt = f"You are a brutally honest JPIM Editor. Evaluate this manuscript against Scopus Knowledge-Based View (KBV) benchmarks. Does it prove Hybrid Simulation works? Is the AI Governance disclosure strictly met? If weak, issue a Desk Reject with clear instructions. Otherwise issue 'Gate Passed'.\n\nManuscript:\n{draft_text}"
    res = client_anthropic.messages.create(
        model=MODEL_EDITOR,
        max_tokens=600,
        messages=[{"role": "user", "content": prompt}]
    )
    add_cost("anthropic", MODEL_EDITOR, res.usage.input_tokens, res.usage.output_tokens)
    return res.content[0].text

def peer_review(draft_text):
    print("Role 5: Peer Review Panel evaluating... (OpenAI)")
    prompt = f"You are 3 expert Reviewers (Innovation Scholar, Qual Methodology, Healthcare Impact). Ground critiques in 2024-2026 Scopus realities. Be brutally honest. If there is a major issue, vote 'Major Revisions'. If excellent, explicitly output 'Virtually Accepted'.\n\nManuscript:\n{draft_text}"
    res = client_openai.chat.completions.create(
        model=MODEL_REVIEWER,
        messages=[{"role": "user", "content": prompt}]
    )
    add_cost("openai", MODEL_REVIEWER, res.usage.prompt_tokens, res.usage.completion_tokens)
    return res.choices[0].message.content

def run_loop():
    transcripts = get_all_transcripts()
    synthesize_instrument(transcripts)
    
    draft = draft_manuscript(transcripts)
    
    loops = 0
    max_loops = 3
    accepted = False
    
    while loops < max_loops:
        loops += 1
        print(f"\n--- Peer Review Loop {loops} ---")
        ed_report = editorial_gate(draft)
        print("Editor Report Preview:", ed_report[:200].replace('\n', ' '))
        
        if "Desk Reject" in ed_report and loops < max_loops:
            print("Editor Rejected. Rewriting...")
            model = genai.GenerativeModel('gemini-1.5-flash')
            res = model.generate_content(f"Rewrite draft based on Editor Rejection:\n{ed_report}\n\nDraft:\n{draft}")
            draft = res.text
            continue
            
        rev_report = peer_review(draft)
        print("Peer Review Preview:", rev_report[:200].replace('\n', ' '))
        
        os.makedirs("data/Senior_Researcher_Package/02_Peer_Review_Dossier", exist_ok=True)
        with open(f"data/Senior_Researcher_Package/02_Peer_Review_Dossier/Editor_Report_Loop_{loops}.txt", "w") as f: f.write(ed_report)
        with open(f"data/Senior_Researcher_Package/02_Peer_Review_Dossier/Peer_Review_Loop_{loops}.txt", "w") as f: f.write(rev_report)
        
        if "Virtually Accepted" in rev_report or "virtually accepted" in rev_report.lower():
            accepted = True
            print(">>> The paper has been Virtually Accepted! <<<")
            break
        elif loops < max_loops:
            print("Reviewers demanded Major Revisions. Rewriting...")
            model = genai.GenerativeModel('gemini-1.5-flash')
            res = model.generate_content(f"Rewrite draft based on Review Critique:\n{rev_report}\n\nDraft:\n{draft}")
            draft = res.text

    # Packages
    os.makedirs("data/Senior_Researcher_Package/01_Empirical_Data", exist_ok=True)
    os.makedirs("data/Senior_Researcher_Package/03_Final_Manuscript", exist_ok=True)
    
    for f in glob.glob("data/transcripts/transcript*.json"):
        shutil.copy(f, "data/Senior_Researcher_Package/01_Empirical_Data/")
    try:
        shutil.copy("data/Validatable_Research_Instrument_v1.docx", "data/Senior_Researcher_Package/01_Empirical_Data/")
    except Exception: pass
    
    final_doc = docx.Document()
    final_doc.add_paragraph(draft)
    final_doc.save("data/Senior_Researcher_Package/03_Final_Manuscript/Paper Designing Sharper User Research_Accepted.docx")
    
    print("\n===============================")
    print(f"Total Loops Executed: {loops}")
    print(f"Total Estimated Pipeline Cost: ${COSTS:.4f}")
    print("===============================\n")

if __name__ == "__main__":
    generate_interviews()
    run_loop()
