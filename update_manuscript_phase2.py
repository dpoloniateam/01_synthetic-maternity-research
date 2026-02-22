import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

source_doc = docx.Document("docs/Repository/Paper Designing Sharper User Research_v0.docx")
new_doc = docx.Document()

# Transfer existing content up to 5. Results
found_results = False
for p in source_doc.paragraphs:
    if "5. Results" in p.text.strip():
        found_results = True
        new_doc.add_paragraph(p.text, style=p.style)
        
        # Inject our new Section 4 (Results) and 5 (Discussion) and the Table
        new_doc.add_heading("Evaluative Analysis of Conversational Yield: From Surface Needs to Latent Insights", level=2)
        new_doc.add_paragraph("A fundamental challenge in the Front End of Innovation (FEI) is transcending surface-level user requests to uncover latent, structural, and emotional drivers that inform breakthrough service design. To evaluate the efficacy of the Synthetic Design Laboratory as a hybrid simulation and high-fidelity micro-sensing capability for innovation, we conducted a comparative analysis of conversational data yielded by distinct prompt calibrations and dynamic model allocations.")
        
        # Add the Comparative Table
        new_doc.add_heading("Methodological Configuration: Dynamic AI Orchestration vs. Static Benchmarks", level=3)
        new_doc.add_paragraph("To prevent provider-level cognitive bias and mimic rigorous peer review, the following multi-agent orchestration architecture was explicitly defined against JPIM constraints:")
        
        table = new_doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Role & Focus'
        hdr_cells[1].text = 'Assigned Provider & Model'
        hdr_cells[2].text = 'Key Features / Performance Modes'
        
        roles = [
            ("Role 1: Interviewer (Analytic Probing)", "OpenAI GPT-5.2 Pro", "Extreme non-linear reasoning; overrides explicit solution generation."),
            ("Role 2: Persona (Empathetic Depth)", "Anthropic Claude Opus 4.6", "Max Context Preservation; sustains 'Private Struggle' narrative organically."),
            ("Role 3: Author (Drafting/Synthesis)", "Google DeepMind Gemini 3.1 Pro", "1M+ token ingestion capacity allowing concurrent document mapping."),
            ("Role 4: JPIM Editor (Gatekeeper)", "Anthropic Claude Sonnet 4.5", "Constitutional Rigor; prevents 'grader-on-its-own-homework' bias."),
            ("Role 5: Review Panel (Multi-Persona)", "OpenAI GPT-4o", "Dynamic Framing; branches context into three isolated expert personalities.")
        ]
        
        for i, (role, model, features) in enumerate(roles):
            row_cells = table.rows[i+1].cells
            row_cells[0].text = role
            row_cells[1].text = model
            row_cells[2].text = features
            
        new_doc.add_paragraph("Our analysis utilized a structured classification rubric that delineates between Needs (i.e., explicit, logistical requirements focusing on 'what' the user wants) and Insights (i.e., deeper understandings revealing hidden truths and structural tensions, focusing on 'why' the user behaves or feels a certain way).")
        
        new_doc.add_heading("Baseline Extraction vs. Calibrated Extraction", level=3)
        new_doc.add_paragraph("In the uncalibrated baseline simulation, the extracted data yielded 100% surface-level Needs, failing to uncover systemic friction. However, following iterative human calibration of the models, the agent produced transcripts composed of 75%-81.25% Insights and only 18.75%-25% Needs. The calibrated synthetic user generated deep, reflective elaboration, yielding powerful latent insights critical for human-centered design in maternity care:")
        
        bullets = [
            ("The 'Stigma of Accommodation' Insight:", " The synthetic participant revealed that requesting logistical support is paradoxically perceived as a capitulation. Utilizing accommodations is internalized not as a right, but as a 'demographic failure' that serves to 'prove them right' regarding her perceived incapacity."),
            ("The 'Private Struggle' Insight:", " The data surfaced a paralyzing latent belief that 'needing help equals failure.' Vulnerable users isolated themselves, hiding severe pain as a behavioral response to mitigate the risk of institutional let-down.")
        ]
        
        for bold_text, normal_text in bullets:
            bp = new_doc.add_paragraph("- ")
            bp.add_run(bold_text).bold = True
            bp.add_run(normal_text)
            
    elif found_results and "6. Conclusions" in p.text.strip():
        # Inject Discussion before Conclusions
        new_doc.add_heading("Discussion: Synthetic Design Laboratory as a Micro-Sensing Capability", level=1)
        new_doc.add_paragraph("The results demonstrate that the Synthetic Design Laboratory operates effectively as a Micro-Sensing Capability at the Front End of Innovation. By functioning as a hybrid simulation environment, it allows organizations to model and surface complex user behaviors—such as the 'Private Struggle'—before entering the field. This capability substantially enhances the Ecological Value of research instruments by ensuring that interview guides are attuned to the actual power dynamics and structural tensions patients face in the real world.")
        new_doc.add_paragraph("In accordance with the Knowledge-Based View (KBV), the iterative calibration of the Generative Interviewer Agent is an explicit knowledge-creation routine. It converts tacit organizational understanding into explicit, high-fidelity research instruments that yield deeper innovation requirements.")
        
        new_doc.add_paragraph(p.text, style=p.style)
        found_results = False
    elif not found_results:
        new_doc.add_paragraph(p.text, style=p.style)

# Appendices
new_doc.add_heading("Appendices", level=1)
new_doc.add_heading("Appendix A: AI Governance and Disclosure", level=2)
new_doc.add_paragraph("In adherence to JPIM guidelines and strict ethical norms, the following elements constituted the AI Governance and Disclosure framework for this study:")
appendix_bullets = [
    ("Model Specification & Targeted Anti-Bias Formulation:", " Simulations securely isolated drafting architectures (Google Gemini) from evaluator architectures (OpenAI & Anthropic) to completely prevent the 'grader-on-its-own-homework' bias during Peer Review testing."),
    ("Human Calibration Routines:", " All underlying prompts framing the AI agents underwent strict human-in-the-loop calibration. Negative constraints overrode the models' default bias towards 'customer service' problem-solving, enforcing an ethnographic methodology instead."),
    ("Audit Trail:", " The system generated a JSON-based Audit Trail recording exact input/output payloads and token costs ensuring full reproducibility."),
    ("Data Privacy:", " AI was restricted entirely to synthetic testbed environments. No actual patient data, or personally identifiable information were processed.")
]
for bold_text, normal_text in appendix_bullets:
    bp = new_doc.add_paragraph("- ")
    bp.add_run(bold_text).bold = True
    bp.add_run(normal_text)

new_doc.save("docs/Repository/Paper Designing Sharper User Research_Draft_For_Review.docx")
print("Saved final manuscript.")
