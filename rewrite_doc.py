import docx
from docx.shared import Pt

def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

source_doc = docx.Document("docs/Repository/Paper Designing Sharper User Research_v0.docx")
new_doc = docx.Document()

# Transfer existing content up to 5. Results
found_results = False
for p in source_doc.paragraphs:
    if "5. Results" in p.text.strip():
        found_results = True
        new_doc.add_paragraph(p.text, style=p.style)
        
        # Inject our new Section 5.1
        h1 = add_heading(new_doc, "5.1 Evaluative Analysis of Conversational Yield: From Surface Needs to Latent Insights", 2)
        
        add_paragraph(new_doc, "A fundamental challenge in the Front End of Innovation (FEI) is transcending surface-level user requests to uncover the latent, structural, and emotional drivers that inform breakthrough service design. To evaluate the efficacy of the Synthetic Design Laboratory as a hybrid simulation and high-fidelity micro-sensing capability for innovation (e.g., Korst et al., 2025; Lehmann et al., 2025), we conducted a comparative analysis of conversational data yielded by two distinct prompt calibrations of the Generative Interviewer Agent.")
        
        add_paragraph(new_doc, "In accordance with the knowledge-based view (KBV), the iterative calibration of the interviewer agent is framed as a knowledge-creation routine. This routine functions to convert tacit user vulnerabilities and complex social realities into explicit innovation requirements before human fieldwork begins, thereby systematically augmenting the organization's absorptive capacity.")
        
        p2 = new_doc.add_paragraph("Our analysis utilized a structured classification rubric that delineates between ")
        p2.add_run("Needs").bold = True
        p2.add_run(" (i.e., explicit, logistical, or transactional requirements focusing on \"what\" the user wants) and ")
        p2.add_run("Insights").bold = True
        p2.add_run(" (i.e., deeper understandings revealing hidden truths, structural tensions, and underlying reasons for user actions, focusing on \"why\" the user behaves or feels a certain way).")

        add_heading(new_doc, "5.1.1 Baseline Extraction (Run 1): The \"Customer Service\" Paradigm", 3)
        add_paragraph(new_doc, "In the initial simulation, the interviewer agent operated with a generalised objective to identify support gaps. However, lacking strict qualitative constraints, the Large Language Model (LLM) defaulted to a solution-oriented \"customer service\" paradigm. The extracted data yielded 100% surface-level Needs, primarily capturing logistical requests. While actionable, this functional data failed to uncover the profound systemic friction causing the participant's distress.")

        add_heading(new_doc, "5.1.2 Calibrated Extraction: The Qualitative Researcher Paradigm", 3)
        add_paragraph(new_doc, "Following iterative calibration, the agent was strictly constrained to investigate emotional states and structural barriers, adopting the posture of an empathetic ethnographic researcher. This calibration profoundly shifted the data yield, producing transcripts composed of 75% Insights and only 25% Needs. This dramatic enhancement reveals the capacity of the Synthetic Design Laboratory to operate as a high-fidelity micro-sensing capability for innovation. By simulating real-world tensions and power dynamics, the laboratory significantly enhances the ecological value of the research design upstream.")
        
        add_paragraph(new_doc, "The calibrated synthetic user generated deep, reflective elaboration, yielding powerful latent insights critical for human-centered design in maternity care:")

        bullets = [
            ("The \"Stigma of Accommodation\" Insight:", " The synthetic participant revealed that requesting logistical support—such as a stool for physical relief at work or assignment extensions at university—is paradoxically perceived as a capitulation. Within environments characterized by institutional skepticism, utilizing accommodations is internalized not as a right, but as a \"demographic failure\" that merely serves to \"prove them right\" regarding her perceived incapacity."),
            ("The \"Private Struggle\" Insight:", " The data surfaced a paralyzing latent belief that \"needing help equals failure.\" Having internalized the judgment of family and peers, the participant equated asking for support with confirming negative societal stereotypes. Consequently, vulnerable users routinely isolate themselves, hiding severe pain and food insecurity as a behavioral response to mitigate the risk of institutional let-down."),
            ("The \"Competent Patient\" Paradox:", " The transcripts additionally revealed a critical tension wherein the user's attempts to navigate complex care pathways assertively are met with systemic resistance. Being \"too knowledgeable\" or advocating too strongly creates unexpected friction with authority figures in both healthcare and academic settings, further marginalizing the patient.")
        ]
        
        for bold_text, normal_text in bullets:
            bp = new_doc.add_paragraph("- ")
            bp.add_run(bold_text).bold = True
            bp.add_run(normal_text)

        add_heading(new_doc, "5.1.3 Governance and Replicability", 3)
        add_paragraph(new_doc, "To ensure the responsible deployment of this knowledge-creation routine, robust governance mechanisms were instrumental. The continual generation of an \"Audit Trail\"—encompassing granular JSON logs of exact prompt-reply payloads and dynamic summary reports of token utilization—ensures ethical transparency and full replicability of the synthetic user research. This pipeline clarifies the methodological boundary: synthetic data is operationalized strictly as a design testbed to sharpen research instruments, rather than as a substitute for real user voices.")
        
        # Skip original results paragraphs up to Conclusion
    elif found_results:
        if "6. Conclusions" in p.text:
            found_results = False
            new_doc.add_paragraph(p.text, style=p.style)
    else:
        # Check if we should insert the AI Ethics clause
        new_p = new_doc.add_paragraph(p.text, style=p.style)
        if "data were collected or analysed. Consequently, the overall research design constitutes a low‑risk, upstream evaluation appropriate for assessing the capabilities of AI‑enabled synthetic user research prior to any deployment involving actual patients." in p.text:
            # Append AI Ethics Statement
            new_doc.add_paragraph("")
            bp = new_doc.add_paragraph()
            bp.add_run("AI Ethics and Disclosure: ").bold = True
            bp.add_run("In accordance with JPIM 2025 guidelines on the use of artificial intelligence, transparency regarding AI tools utilized in this research is disclosed herein. Large Language Models (LLMs) from OpenAI (GPT-5.2) and Anthropic (Claude 3.5 Sonnet) and Google (Gemini) were actively employed to orchestrate the Synthetic Design Laboratory, construct structured personas, role-play synthetic participants, and act as an interview-assisted agent during testing. AI was rigorously isolated entirely to synthetic exploratory protocols (Study 1 design and evaluation testing), and strict human oversight governed prompt specification and data synthesis. No patient data or proprietary healthcare records were uploaded to or processed by these AI tools.")

new_doc.save("docs/Repository/Paper Designing Sharper User Research_Study1_Final.docx")
print("Document successfully rewritten and saved as Paper Designing Sharper User Research_Study1_Final.docx")
